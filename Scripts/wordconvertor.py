import os
import shutil
from docx import Document
from deep_translator.exceptions import LanguageNotSupportedException
from deep_translator import GoogleTranslator

def copy_translate_docx():
    # Get input file from user
    source_path = input("Enter the path to the input .docx file: ")
    if not os.path.isfile(source_path) or not source_path.lower().endswith('.docx'):
        print("Error: Input file not found or not a .docx file.")
        return
    
    # Get destination directory from user
    dest_dir = input("Enter the destination directory: ")
    if not os.path.isdir(dest_dir):
        print("Error: Destination directory not found.")
        return
    
    # Get output language from user
    output_lang = input("Enter the output language (e.g., 'es' for Spanish): ")
    
    try:
        # Check if the output language is supported by Google Translate
        GoogleTranslator.validate_language(output_lang)
    except LanguageNotSupportedException:
        print(f"Error: Language '{output_lang}' is not currently supported by Google Translate.")
        return
    
    # Copy the .docx file
    source_name = os.path.basename(source_path)
    dest_name = source_name.replace('.docx', '_copy.docx')
    dest_path = os.path.join(dest_dir, dest_name)
    shutil.copyfile(source_path, dest_path)

    # Open the copied .docx file
    doc = Document(dest_path)
    
    translated_doc = Document()
    
    # Translate each paragraph
    for para in doc.paragraphs:
        translated_para = translated_doc.add_paragraph()
        
        # Translate each run in the paragraph
        for run in para.runs:
            translated_text = GoogleTranslator(source='auto', target=output_lang).translate(run.text)
            translated_run = translated_para.add_run(translated_text)
            
            # Preserve the original style of the run
            translated_run.bold = run.bold
            translated_run.italic = run.italic
            translated_run.underline = run.underline
            translated_run.font.size = run.font.size
            translated_run.font.name = run.font.name
            translated_run.font.color.rgb = run.font.color.rgb
    
    # Save the translated .docx file
    translated_doc.save(dest_path)

    print("Translation complete. Translated file saved at:", dest_path)

# Example usage:
copy_translate_docx()